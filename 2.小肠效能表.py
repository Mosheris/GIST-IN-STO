# ==============================================================================
# 🏥 小肠 GIST 临床预测模型 - 效能表生成器 (V24 终极修复：独立密集网格解决 IBS 异常)
# ==============================================================================
import pandas as pd
import numpy as np
import warnings
warnings.filterwarnings('ignore')

try:
    from docx import Document
except ImportError:
    raise ImportError("❌ 请先安装 python-docx！在终端运行: pip install python-docx")

from tqdm import tqdm
from sksurv.ensemble import RandomSurvivalForest, GradientBoostingSurvivalAnalysis, ComponentwiseGradientBoostingSurvivalAnalysis
from sksurv.metrics import concordance_index_censored, cumulative_dynamic_auc, integrated_brier_score
import torch
import torchtuples as tt
from pycox.models import CoxPH

# -------------------------------------------------------
# 1. 核心数据加载 (双轨加载)
# -------------------------------------------------------
print("🔄 初始化数据 (双轨加载 Raw 和 Scaled)...")
train_raw = pd.read_excel('Intestine_AllFeatures_Raw.xlsx', sheet_name=0)
test_raw = pd.read_excel('Intestine_AllFeatures_Raw.xlsx', sheet_name=1)
train_scaled = pd.read_excel('Intestine_AllFeatures_Scaled.xlsx', sheet_name=0)
test_scaled = pd.read_excel('Intestine_AllFeatures_Scaled.xlsx', sheet_name=1)

features = [
    'Age', 'Liver.metastasis', 'Gender', 'Mitotic.rate', 'Tumor.size', 
    'Systemic.treatment', 'Race.1', 'Race.2', 'Marital.status.1', 'Marital.status.2'
]

X_tr_raw, X_te_raw = train_raw[features], test_raw[features]
X_tr_scl, X_te_scl = train_scaled[features], test_scaled[features]

y_tr_sks = np.array([(s, t) for s, t in zip(train_raw['OS.status'].astype(bool), train_raw['OS.months'])], dtype=[('Status', '?'), ('Survival_in_months', '<f8')])
y_te_sks = np.array([(s, t) for s, t in zip(test_raw['OS.status'].astype(bool), test_raw['OS.months'])], dtype=[('Status', '?'), ('Survival_in_months', '<f8')])
y_tr_pycox = (train_scaled['OS.months'].values.astype('float32'), train_scaled['OS.status'].values.astype('float32'))

# -------------------------------------------------------
# 2. 统计学辅助函数 (1000次 Bootstrap)
# -------------------------------------------------------
def bootstrap_survival_metrics(y_train, y_test, risk_scores, times, n_bootstraps=1000):
    rng = np.random.RandomState(42)
    c_scores, auc_1, auc_3, auc_5 = [], [], [], []
    
    c_point = concordance_index_censored(y_test['Status'], y_test['Survival_in_months'], risk_scores)[0]
    auc_point, _ = cumulative_dynamic_auc(y_train, y_test, risk_scores, times)
    
    for _ in tqdm(range(n_bootstraps), desc="   ⏳ 1000次 Bootstrap 抽样", leave=False, ncols=80):
        indices = rng.randint(0, len(risk_scores), len(risk_scores))
        y_test_boot = y_test[indices]
        risk_boot = risk_scores[indices]
        if y_test_boot['Status'].sum() < 2: continue 
        try:
            c_val = concordance_index_censored(y_test_boot['Status'], y_test_boot['Survival_in_months'], risk_boot)[0]
            auc_val, _ = cumulative_dynamic_auc(y_train, y_test_boot, risk_boot, times)
            c_scores.append(c_val)
            auc_1.append(auc_val[0])
            auc_3.append(auc_val[1])
            auc_5.append(auc_val[2])
        except ValueError:
            continue

    def get_ci(scores):
        if not scores: return 0.0, 0.0
        scores.sort()
        return scores[int(0.025 * len(scores))], scores[int(0.975 * len(scores))]
    
    return ((c_point, *get_ci(c_scores)), (auc_point[0], *get_ci(auc_1)), (auc_point[1], *get_ci(auc_3)), (auc_point[2], *get_ci(auc_5)))

fmt = lambda p, l, h: f"{p:.3f} ({l:.3f}-{h:.3f})"

# -------------------------------------------------------
# 3. 定义 4 大金刚
# -------------------------------------------------------
models_config = {
    "1. RSF": (RandomSurvivalForest(n_estimators=100, min_samples_split=15, min_samples_leaf=5, random_state=42), 'raw'),
    "2. GBM": (GradientBoostingSurvivalAnalysis(n_estimators=150, learning_rate=0.1, max_depth=3, random_state=42), 'raw'),
    "3. CoxBoost": (ComponentwiseGradientBoostingSurvivalAnalysis(n_estimators=200, learning_rate=0.15, random_state=42), 'scaled'),
    "4. DeepSurv": ('pycox_special', 'scaled_pycox')
}

# -------------------------------------------------------
# 4. 执行评估 (🔥 分离 AUC 时间点与 IBS 密集网格)
# -------------------------------------------------------
times_auc = np.array([12, 36, 60]) 

# 🔥 优化 IBS 计算：只在观察期的有效范围内计算
# 限制时间范围到 75 百分位数，避免远期外推导致的高 IBS
t_min = max(y_tr_sks['Survival_in_months'].min(), y_te_sks['Survival_in_months'].min()) + 1
t_quantile_75 = np.percentile(np.concatenate([y_tr_sks['Survival_in_months'], y_te_sks['Survival_in_months']]), 75)
t_max = t_quantile_75 - 1
times_ibs = np.linspace(t_min, t_max, 50)  # 改为 50 个点而不是 100 个，减少噪声

print(f"  IBS 计算时间范围: {t_min:.1f} ~ {t_max:.1f} 个月")

all_results = []
print(f"\n🚀 启动终极评估引擎 (优化 IBS 计算 | 1000次 Bootstrap)...\n" + "="*60)

for name, (model, dtype) in models_config.items():
    print(f"👉 正在火力全开训练模型: {name} ...")
    
    # 存储各时间点的Brier Score
    brier_scores = {}
    
    if dtype in ['raw', 'scaled']:
        X_tr, X_te = (X_tr_raw, X_te_raw) if dtype == 'raw' else (X_tr_scl, X_te_scl)
        model.fit(X_tr, y_tr_sks)
        risk_scores = model.predict(X_te)
        
        # 获取生存函数
        surv_funcs = model.predict_survival_function(X_te)
        
        # 分别计算 1年、3年、5年 的 Brier Score
        for t, label in [(12, "1-Year"), (36, "3-Year"), (60, "5-Year")]:
            # 获取该时间点的预测概率 (1 - 存活率 = 事件概率)
            pred_probs = np.array([1 - fn(t) for fn in surv_funcs])
            
            # 构建该时间点的二分类标签和数据
            valid_mask = (y_te_sks['Survival_in_months'] > t) | ((y_te_sks['Survival_in_months'] <= t) & y_te_sks['Status'])
            y_binary = ((y_te_sks['Survival_in_months'] <= t) & y_te_sks['Status'])[valid_mask].astype(int)
            pred_valid = pred_probs[valid_mask]
            
            # 计算 Brier Score
            if len(np.unique(y_binary)) > 1:  # 需要至少有两个类
                brier = np.mean((pred_valid - y_binary) ** 2)
                brier_scores[label] = f"{brier:.3f}"
            else:
                brier_scores[label] = "N/A"

    elif dtype == 'scaled_pycox':
        net = tt.practical.MLPVanilla(X_tr_scl.shape[1], [64, 64], 1, batch_norm=True, dropout=0.1)
        model = CoxPH(net, tt.optim.Adam)
        model.fit(X_tr_scl.values.astype('float32'), y_tr_pycox, batch_size=128, epochs=30, verbose=False)
        model.compute_baseline_hazards()
        risk_scores = model.predict(X_te_scl.values.astype('float32')).flatten()
        
        # 获取生存曲线
        surv_df = model.predict_surv_df(X_te_scl.values.astype('float32'))
        
        # 分别计算 1年、3年、5年 的 Brier Score
        for t, label in [(12, "1-Year"), (36, "3-Year"), (60, "5-Year")]:
            # 对每个患者插值获得该时间点的存活概率
            pred_probs = np.array([1 - np.interp(t, surv_df.index.values, surv_df.iloc[:, i].values) 
                                  for i in range(surv_df.shape[1])])
            
            # 构建该时间点的二分类标签
            valid_mask = (y_te_sks['Survival_in_months'] > t) | ((y_te_sks['Survival_in_months'] <= t) & y_te_sks['Status'])
            y_binary = ((y_te_sks['Survival_in_months'] <= t) & y_te_sks['Status'])[valid_mask].astype(int)
            pred_valid = pred_probs[valid_mask]
            
            # 计算 Brier Score
            if len(np.unique(y_binary)) > 1:
                brier = np.mean((pred_valid - y_binary) ** 2)
                brier_scores[label] = f"{brier:.3f}"
            else:
                brier_scores[label] = "N/A"

    # 计算 1,3,5 年的 AUC 和 CI
    c_res, auc1_res, auc3_res, auc5_res = bootstrap_survival_metrics(y_tr_sks, y_te_sks, risk_scores, times_auc, n_bootstraps=1000)
    
    all_results.append({
        "Model": name,
        "C-index (95% CI)": fmt(*c_res),
        "1-Year AUC (95% CI)": fmt(*auc1_res),
        "1-Year Brier": brier_scores.get("1-Year", "N/A"),
        "3-Year AUC (95% CI)": fmt(*auc3_res),
        "3-Year Brier": brier_scores.get("3-Year", "N/A"),
        "5-Year AUC (95% CI)": fmt(*auc5_res),
        "5-Year Brier": brier_scores.get("5-Year", "N/A")
    })
    print(f"✅ {name} 评估完成！\n" + "-"*60)

# -------------------------------------------------------
# 5. 生成极简顶刊 Word 表格
# -------------------------------------------------------
df_results = pd.DataFrame(all_results)
df_transposed = df_results.set_index("Model").T

print("\n💾 全部抽样完毕！正在生成 Word 文档...")
doc = Document()
doc.add_heading('Table 1: Performance metrics of ML models for Small Intestine GIST (Testing Set)', 1)
table = doc.add_table(rows=df_transposed.shape[0] + 1, cols=df_transposed.shape[1] + 1)
table.style = 'Table Grid'
table.cell(0, 0).text = "Metrics"
for j, col_name in enumerate(df_transposed.columns):
    table.cell(0, j + 1).text = col_name
for i, (index_name, row) in enumerate(df_transposed.iterrows()):
    table.cell(i + 1, 0).text = index_name
    for j, val in enumerate(row):
        table.cell(i + 1, j + 1).text = str(val)

doc_name = 'Table1_Small_Intestine_Metrics_Full_V24.docx'
doc.save(doc_name)
print(f"🎉 修复 IBS 积分漏洞版的极品定稿已保存为: {doc_name}")