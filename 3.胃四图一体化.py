# ==============================================================================
# 🏥 小肠 GIST 临床预测模型 - V27 究极真理版 (表图绝对同源 + 解锁动态 AUC)
# ==============================================================================
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import warnings
from tqdm import tqdm

try:
    from docx import Document
except ImportError:
    raise ImportError("❌ 请先安装 python-docx！在终端运行: pip install python-docx")

from sksurv.ensemble import RandomSurvivalForest, GradientBoostingSurvivalAnalysis, ComponentwiseGradientBoostingSurvivalAnalysis
from sksurv.metrics import concordance_index_censored, cumulative_dynamic_auc
from sksurv.nonparametric import kaplan_meier_estimator
from sklearn.metrics import roc_curve, precision_recall_curve, average_precision_score
from sklearn.calibration import calibration_curve
import torch
import torchtuples as tt
from pycox.models import CoxPH

warnings.filterwarnings('ignore')

# 🔥 锁死随机种子，杜绝 DeepSurv 波动
np.random.seed(42)
torch.manual_seed(42)

plt.style.use('seaborn-v0_8-whitegrid')
plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False  
COLORS = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728']

# -------------------------------------------------------
# 1. 核心数据加载
# -------------------------------------------------------
print("🔄 [1/5] 初始化数据...")
train_raw = pd.read_excel('Stomach_AllFeatures_Raw.xlsx', sheet_name=0)
test_raw = pd.read_excel('Stomach_AllFeatures_Raw.xlsx', sheet_name=1)
train_scaled = pd.read_excel('Stomach_AllFeatures_Scaled.xlsx', sheet_name=0)
test_scaled = pd.read_excel('Stomach_AllFeatures_Scaled.xlsx', sheet_name=1)

features = ['Age', 'Liver.metastasis', 'Gender', 'Mitotic.rate', 'Tumor.size', 
            'Race.1', 'Race.2', 'Marital.status.1', 'Marital.status.2']

X_tr_raw, X_te_raw = train_raw[features], test_raw[features]
X_tr_scl, X_te_scl = train_scaled[features], test_scaled[features]

y_tr_sks = np.array([(s, t) for s, t in zip(train_raw['OS.status'].astype(bool), train_raw['OS.months'])], dtype=[('Status', '?'), ('Survival_in_months', '<f8')])
y_te_sks = np.array([(s, t) for s, t in zip(test_raw['OS.status'].astype(bool), test_raw['OS.months'])], dtype=[('Status', '?'), ('Survival_in_months', '<f8')])
y_tr_pycox = (train_scaled['OS.months'].values.astype('float32'), train_scaled['OS.status'].values.astype('float32'))

# -------------------------------------------------------
# 2. 手写底层核心：IPCW 加权 (彻底解决图表差异)
# -------------------------------------------------------
def get_ipcw_weights(y_train, y_test, t):
    censor_status = ~y_train['Status']
    times_km, G_prob = kaplan_meier_estimator(censor_status, y_train['Survival_in_months'])
    
    def get_G(t_val):
        idx = np.searchsorted(times_km, t_val, side='right') - 1
        return G_prob[idx] if idx >= 0 and G_prob[idx] > 0 else 1e-5
    
    weights, y_binary = np.zeros(len(y_test)), np.zeros(len(y_test))
    valid_mask = np.ones(len(y_test), dtype=bool)
    G_t = get_G(t)
    
    for i, (status, time_i) in enumerate(zip(y_test['Status'], y_test['Survival_in_months'])):
        if time_i <= t and status:
            weights[i] = 1.0 / get_G(time_i)
            y_binary[i] = 1
        elif time_i > t:
            weights[i] = 1.0 / G_t
            y_binary[i] = 0
        else:
            valid_mask[i] = False
            
    return valid_mask, y_binary, weights

# -------------------------------------------------------
# 3. 模型训练与动态概率提取
# -------------------------------------------------------
print("🚀 [2/5] 正在训练 4 大金刚并提取时间依赖性概率...")
times = [12, 36, 60]
time_labels = ['1-Year', '3-Year', '5-Year']

models_config = {
    "1. RSF": (RandomSurvivalForest(n_estimators=150, min_samples_split=15, min_samples_leaf=5, random_state=42), 'raw'),
    "2. GBM": (GradientBoostingSurvivalAnalysis(n_estimators=150, learning_rate=0.15, max_depth=3, random_state=42), 'raw'),
    "3. CoxBoost": (ComponentwiseGradientBoostingSurvivalAnalysis(n_estimators=200, learning_rate=0.15, random_state=42), 'scaled'),
    "4. DeepSurv": ('pycox_special', 'scaled_pycox')
}

model_risk_scores = {} # 供 C-index 评判总排序
model_surv_probs = {}  # 供 AUC/ROC 评判动态预测

for name, (model, dtype) in models_config.items():
    if dtype in ['raw', 'scaled']:
        X_tr, X_te = (X_tr_raw, X_te_raw) if dtype == 'raw' else (X_tr_scl, X_te_scl)
        model.fit(X_tr, y_tr_sks)
        model_risk_scores[name] = model.predict(X_te)
        surv_funcs = model.predict_survival_function(X_te)
        model_surv_probs[name] = {label: np.array([1 - fn(t) for fn in surv_funcs]) for t, label in zip(times, time_labels)}
        
    elif dtype == 'scaled_pycox':
        net = tt.practical.MLPVanilla(X_tr_scl.shape[1], [32, 32], 1, batch_norm=True, dropout=0.1)
        model = CoxPH(net, tt.optim.Adam)
        model.fit(X_tr_scl.values.astype('float32'), y_tr_pycox, batch_size=32, epochs=30, verbose=False)
        model.compute_baseline_hazards()
        model_risk_scores[name] = model.predict(X_te_scl.values.astype('float32')).flatten()
        surv_df = model.predict_surv_df(X_te_scl.values.astype('float32'))
        model_surv_probs[name] = {label: np.array([1 - np.interp(t, surv_df.index.values, surv_df.iloc[:, i].values) for i in range(surv_df.shape[1])]) for t, label in zip(times, time_labels)}

# -------------------------------------------------------
# 4. 执行 IPCW 1000次抽样计算 (一次计算，表图通用)
# -------------------------------------------------------
print("⏳ [3/5] 正在执行核心 IPCW Bootstrap 抽样引擎...")
model_metrics = {}
plot_roc_data = {name: {} for name in models_config.keys()}
fmt = lambda p, l, h: f"{p:.3f} ({l:.3f}-{h:.3f})"
times_auc = np.array(times)

for name in models_config.keys():
    metrics_dict = {}
    risk_scores = model_risk_scores[name]
    auc_point, _ = cumulative_dynamic_auc(y_tr_sks, y_te_sks, risk_scores, times_auc)
    
    # 1. 计算 C-index (全局排序)
    c_base = concordance_index_censored(y_te_sks['Status'], y_te_sks['Survival_in_months'], risk_scores)[0]
    c_boot = []
    rng = np.random.RandomState(42)
    for _ in range(1000):
        idx = rng.randint(0, len(y_te_sks), len(y_te_sks))
        y_b, r_b = y_te_sks[idx], risk_scores[idx]
        if y_b['Status'].sum() > 1:
            try: c_boot.append(concordance_index_censored(y_b['Status'], y_b['Survival_in_months'], r_b)[0])
            except: pass
    c_boot.sort()
    metrics_dict["C-index (95% CI)"] = fmt(c_base, c_boot[25], c_boot[975])

    # 2. 计算动态时间点 AUC、Brier 与 画图数据
    for t_idx, (t, label) in enumerate(zip(times, time_labels)):
        valid_mask, y_binary, weights = get_ipcw_weights(y_tr_sks, y_te_sks, t)
        y_p = model_surv_probs[name][label][valid_mask]
        y_score = risk_scores[valid_mask]
        y_b = y_binary[valid_mask]
        w = weights[valid_mask]
        
        # 绝对精准的基础 AUC 和 Brier
        base_auc = auc_point[t_idx]
        base_brier = np.average((y_p - y_b)**2, weights=w) if len(np.unique(y_b)) > 1 else np.nan
        fpr_base, tpr_base, _ = roc_curve(y_b, y_score, sample_weight=w)
        
        # 1000 次 Bootstrap
        auc_boot, tprs_boot = [], []
        mean_fpr = np.linspace(0, 1, 100)
        
        for _ in tqdm(range(1000), desc=f"   {name} {label}", leave=False, ncols=80):
            idx = rng.randint(0, len(y_b), len(y_b))
            y_b_i, y_s_i, w_i = y_b[idx], y_score[idx], w[idx]
            if len(np.unique(y_b_i)) < 2: continue
            auc_val, _ = cumulative_dynamic_auc(y_tr_sks, y_te_sks[idx], risk_scores[idx], np.array([t]))
            auc_boot.append(auc_val[0])
            fpr_i, tpr_i, _ = roc_curve(y_b_i, y_s_i, sample_weight=w_i)
            tpr_interp = np.interp(mean_fpr, fpr_i, tpr_i)
            tpr_interp[0] = 0.0
            tprs_boot.append(tpr_interp)
            
        auc_boot.sort()
        auc_l, auc_h = auc_boot[25], auc_boot[975]
        metrics_dict[f"{label} AUC (95% CI)"] = fmt(base_auc, auc_l, auc_h)
        metrics_dict[f"{label} Brier"] = f"{base_brier:.3f}" if not np.isnan(base_brier) else "N/A"
        
        # 存储极其完美的画图数据
        tpr_lower = np.percentile(tprs_boot, 2.5, axis=0)
        tpr_upper = np.percentile(tprs_boot, 97.5, axis=0)
        plot_roc_data[name][label] = (fpr_base, tpr_base, mean_fpr, tpr_lower, tpr_upper, base_auc, auc_l, auc_h)

    model_metrics[name] = metrics_dict

# -------------------------------------------------------
# 5. 生成极其同步的 Word 表格
# -------------------------------------------------------
print("💾 [4/5] 正在生成金标准 Word 表格...")
all_results = [{"Model": name, **metrics} for name, metrics in model_metrics.items()]
df_transposed = pd.DataFrame(all_results).set_index("Model").T

doc = Document()
doc.add_heading('Table 1: Performance metrics of ML models for Small Intestine GIST', 1)
table = doc.add_table(rows=df_transposed.shape[0] + 1, cols=df_transposed.shape[1] + 1)
table.style = 'Table Grid'
table.cell(0, 0).text = "Metrics"
for j, col_name in enumerate(df_transposed.columns): table.cell(0, j + 1).text = col_name
for i, (index_name, row) in enumerate(df_transposed.iterrows()):
    table.cell(i + 1, 0).text = index_name
    for j, val in enumerate(row): table.cell(i + 1, j + 1).text = str(val)
doc.save('Table1_Metrics_Ultimate_V27.docx')

# -------------------------------------------------------
# 6. 生成彻底同步的高清效能图
# -------------------------------------------------------
print("🎨 [5/5] 正在绘制顶刊级别效能图 (图表数值 100% 绝对锚定)...")

# --- 图 A: ROC ---
fig, axes = plt.subplots(1, 3, figsize=(24, 8))
for ax, t, label in zip(axes, times, time_labels):
    for i, name in enumerate(models_config.keys()):
        fpr_base, tpr_base, mean_fpr, l_tpr, u_tpr, auc_point, auc_l, auc_h = plot_roc_data[name][label]
        legend_label = f"{name} (AUC={auc_point:.3f}, 95% CI: {auc_l:.3f}-{auc_h:.3f})"
        ax.plot(fpr_base, tpr_base, color=COLORS[i], label=legend_label, lw=2.5)
        ax.fill_between(mean_fpr, l_tpr, u_tpr, color=COLORS[i], alpha=0.15)
        
    ax.plot([0, 1], [0, 1], 'k--', lw=1.5, alpha=0.6)
    ax.set_xlabel('False Positive Rate', fontsize=14, fontweight='bold')
    ax.set_ylabel('True Positive Rate', fontsize=14, fontweight='bold')
    ax.set_title(f'{label} ROC Curve', fontsize=16, fontweight='bold')
    ax.legend(loc="lower right", fontsize=11, frameon=True, edgecolor='gray')
    ax.grid(True, linestyle='--', alpha=0.4); ax.set_xlim([-0.02, 1.02]); ax.set_ylim([-0.02, 1.02])
plt.tight_layout(); plt.savefig('Fig1_ROC_Ultimate_V27.pdf')

# --- 图 B: PR 曲线 ---
fig, axes = plt.subplots(1, 3, figsize=(24, 8))
for ax, t, label in zip(axes, times, time_labels):
    valid_mask, y_binary, weights = get_ipcw_weights(y_tr_sks, y_te_sks, t)
    y_b, w = y_binary[valid_mask], weights[valid_mask]
    baseline = np.sum(y_b * w) / np.sum(w)
    ax.plot([0, 1], [baseline, baseline], 'k--', lw=2, label=f'Baseline (AP={baseline:.3f})')
    
    for i, name in enumerate(models_config.keys()):
        y_p = model_surv_probs[name][label][valid_mask]
        precision, recall, _ = precision_recall_curve(y_b, y_p, sample_weight=w)
        ap = average_precision_score(y_b, y_p, sample_weight=w)
        ax.plot(recall, precision, lw=2.5, color=COLORS[i], alpha=0.9, label=f'{name} (AP={ap:.3f})')
        
    ax.set_xlabel('Recall', fontsize=14, fontweight='bold'); ax.set_ylabel('Precision', fontsize=14, fontweight='bold')
    ax.set_title(f'{label} PR Curve', fontsize=16, fontweight='bold')
    ax.legend(loc='upper right', fontsize=11, frameon=True)
    ax.grid(True, linestyle='--', alpha=0.5); ax.set_xlim([0, 1]); ax.set_ylim([0, 1.02])
plt.tight_layout(); plt.savefig('Fig2_PR_Ultimate_V27.pdf')

# --- 图 C: 校准曲线 ---
fig, axes = plt.subplots(1, 3, figsize=(24, 8))
for ax, t, label in zip(axes, times, time_labels):
    ax.plot([0, 1], [0, 1], "k:", lw=2, label="Perfectly Calibrated")
    valid_mask, y_binary, weights = get_ipcw_weights(y_tr_sks, y_te_sks, t)
    for i, name in enumerate(models_config.keys()):
        y_p = model_surv_probs[name][label][valid_mask]
        prob_true, prob_pred = calibration_curve(y_binary[valid_mask], y_p, n_bins=6, strategy='quantile')
        ax.plot(prob_pred, prob_true, "s-", lw=2.5, markersize=8, color=COLORS[i], alpha=0.9, label=name)
        
    ax.set_xlabel(f'Mean Predicted {label} Mortality', fontsize=14, fontweight='bold'); ax.set_ylabel('Observed Mortality', fontsize=14, fontweight='bold')
    ax.set_title(f'{label} Calibration Curve', fontsize=16, fontweight='bold')
    ax.legend(loc='lower right', fontsize=11, frameon=True)
    ax.grid(True, linestyle='--', alpha=0.5); ax.set_xlim([0, 1]); ax.set_ylim([0, 1])
plt.tight_layout(); plt.savefig('Fig3_Calibration_Ultimate_V27.pdf')

# --- 图 D: DCA 曲线 ---
fig, axes = plt.subplots(1, 3, figsize=(24, 8))
for ax, t, label in zip(axes, times, time_labels):
    thresholds = np.linspace(0.01, 0.99, 100)
    ax.plot([0, 1], [0, 0], 'k-', lw=2, label='Treat None')
    
    valid_mask, y_binary, weights = get_ipcw_weights(y_tr_sks, y_te_sks, t)
    y_b, w = y_binary[valid_mask], weights[valid_mask]
    total_w = np.sum(w)
    prevalence = np.sum(y_b * w) / total_w
    
    treat_all = [prevalence - (1 - prevalence) * (th / (1 - th)) if th < 1 else -0.1 for th in thresholds]
    ax.plot(thresholds, treat_all, 'gray', linestyle=':', lw=2, label='Treat All')
    
    for i, name in enumerate(models_config.keys()):
        y_p = model_surv_probs[name][label][valid_mask]
        net_benefits = []
        for th in thresholds:
            pred = (y_p >= th).astype(int)
            tp, fp = np.sum(w * ((y_b == 1) & (pred == 1))), np.sum(w * ((y_b == 0) & (pred == 1)))
            net_benefits.append((tp / total_w) - (fp / total_w) * (th / (1 - th)) if th < 1 else -0.1)
        ax.plot(thresholds, net_benefits, lw=2.5, color=COLORS[i], alpha=0.9, label=name)
        
    ax.set_xlabel('Threshold Probability', fontsize=14, fontweight='bold'); ax.set_ylabel('Net Benefit', fontsize=14, fontweight='bold')
    ax.set_title(f'{label} DCA', fontsize=16, fontweight='bold')
    ax.legend(loc='upper right', fontsize=11, frameon=True)
    ax.grid(True, linestyle='--', alpha=0.5); ax.set_xlim([0, 1]); ax.set_ylim([-0.1, max(prevalence+0.1, 0.4)])
plt.tight_layout(); plt.savefig('Fig4_DCA_Ultimate_V27.pdf')

print("\n🎉🎉🎉 战役彻底结束！图表数据底层同源，完美同步，尽情享受顶级出图效果吧！") 